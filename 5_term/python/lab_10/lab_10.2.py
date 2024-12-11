from docx import Document
from docx.shared import Inches

doc = Document('C:\\Users\\Бобр Даша\\Desktop\\university\\3 КУРС\\5 сем\\Язык Python для работы с данными\\lab_10\\lab_10.docx')

img_path = 'C:\\Users\\Бобр Даша\\Desktop\\university\\3 КУРС\\5 сем\\Язык Python для работы с данными\\lab_10\\photo.jpg'  
doc.add_picture(img_path, width=Inches(2)) 

doc.add_paragraph('Милый котик')  

doc.save('C:\\Users\\Бобр Даша\\Desktop\\university\\3 КУРС\\5 сем\\Язык Python для работы с данными\\lab_10\\lab_10.docx')
