from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_paragraph('В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:\n')

doc.add_paragraph('○ Флеш-память: используется для хранения скетчей.').paragraph_format.left_indent = Pt(20)
p = doc.add_paragraph()
p.add_run('○ ОЗУ(')
p.add_run('SRAM').bold = True
p.add_run(' — ')
p.add_run('static random access memory').italic = True
p.add_run(', статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.')
p.paragraph_format.left_indent = Pt(20)
doc.add_paragraph('○ EEPROM (энергонезависимая память): используется для хранения постоянной информации.').paragraph_format.left_indent = Pt(20)

doc.add_paragraph("Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.\n")


table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'

tab_data = [
    ['', 'ATmega168', 'ATmega328', 'ATmega1280', 'ATmega2560'],
    ['Flash (1 кБ flash-памяти занят загрузчиком)', '16 Кбайт', '32 Кбайт', '128 Кбайт', '256 Кбайт'],
    ['SRAM', '1 Кбайт', '2 Кбайт', '8 байта', '8 Кбайт'],
    ['EEPROM', '512 Кбайт', '1024 Кбайт', '4 Кбайт', '4 Кбайт'],
]

for i in range(len(tab_data)):
    for j in range(len(tab_data[i])):
        cell = table.cell(i, j)
        cell.text = tab_data[i][j]
        
        if i == 0 or j == 0:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True  

text = (
    "\nПамять EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°С. Эти данные не распространяются на операции чтения данных из EEPROM - чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM."
)

p = doc.add_paragraph().add_run(text).italic = True  

doc.save('C:\\Users\\Бобр Даша\\Desktop\\university\\3 КУРС\\5 сем\\Язык Python для работы с данными\\lab_10\\lab_10.docx')
